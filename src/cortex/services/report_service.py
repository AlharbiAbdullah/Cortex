"""
Automated Report Generation Service.

Generates formatted reports from data and templates in multiple formats
including PDF, DOCX, and HTML.
"""

import base64
import io
import logging
import os
from datetime import datetime
from functools import lru_cache
from pathlib import Path
from typing import Any

from jinja2 import Environment, FileSystemLoader, select_autoescape
from pydantic import BaseModel, Field
from pydantic_settings import BaseSettings, SettingsConfigDict

logger = logging.getLogger(__name__)


class ReportSettings(BaseSettings):
    """Report service configuration."""

    model_config = SettingsConfigDict(env_file=".env", extra="ignore")

    template_dir: str = "templates/reports"
    output_dir: str = "generated_files/reports"
    default_format: str = "html"
    company_name: str = "Cortex"
    logo_path: str | None = None


@lru_cache()
def get_report_settings() -> ReportSettings:
    """Get cached report settings."""
    return ReportSettings()


class ChartConfig(BaseModel):
    """Configuration for embedding charts in reports."""

    chart_type: str = Field(
        ...,
        description="Chart type: bar, line, pie, scatter",
    )
    title: str = Field(..., description="Chart title")
    data: dict[str, Any] = Field(..., description="Chart data")
    width: int = Field(default=600, description="Chart width in pixels")
    height: int = Field(default=400, description="Chart height in pixels")


class ReportTemplate(BaseModel):
    """Report template definition."""

    template_id: str = Field(..., description="Unique template identifier")
    name: str = Field(..., description="Template display name")
    description: str = Field(default="", description="Template description")
    template_content: str = Field(..., description="Jinja2 template content")
    required_fields: list[str] = Field(
        default_factory=list,
        description="Required data fields",
    )
    sample_data: dict[str, Any] = Field(
        default_factory=dict,
        description="Sample data for preview",
    )


class GeneratedReport(BaseModel):
    """Generated report metadata."""

    report_id: str = Field(..., description="Unique report identifier")
    template_id: str = Field(..., description="Template used")
    format: str = Field(..., description="Output format")
    filename: str = Field(..., description="Generated filename")
    file_path: str = Field(..., description="File path")
    file_size: int = Field(default=0, description="File size in bytes")
    generated_at: str = Field(..., description="Generation timestamp")
    download_url: str = Field(default="", description="Download URL")


class ReportService:
    """
    Service for generating automated reports.

    Supports:
    - Template-based generation using Jinja2
    - Chart embedding (matplotlib/plotly)
    - Multi-format output (PDF, DOCX, HTML)
    - Custom branding
    """

    def __init__(self):
        """Initialize the report service."""
        self._settings = get_report_settings()
        self._ensure_directories()
        self._setup_jinja_env()
        self._builtin_templates = self._load_builtin_templates()

    def _ensure_directories(self) -> None:
        """Ensure required directories exist."""
        Path(self._settings.output_dir).mkdir(parents=True, exist_ok=True)
        Path(self._settings.template_dir).mkdir(parents=True, exist_ok=True)

    def _setup_jinja_env(self) -> None:
        """Set up Jinja2 environment."""
        template_path = Path(self._settings.template_dir)

        if template_path.exists():
            self._jinja_env = Environment(
                loader=FileSystemLoader(str(template_path)),
                autoescape=select_autoescape(["html", "xml"]),
            )
        else:
            # Use string-based templates
            self._jinja_env = Environment(
                autoescape=select_autoescape(["html", "xml"]),
            )

        # Add custom filters
        self._jinja_env.filters["datetime"] = self._format_datetime
        self._jinja_env.filters["currency"] = self._format_currency
        self._jinja_env.filters["percentage"] = self._format_percentage

    def _load_builtin_templates(self) -> dict[str, ReportTemplate]:
        """Load built-in report templates."""
        return {
            "executive_summary": ReportTemplate(
                template_id="executive_summary",
                name="Executive Summary Report",
                description="High-level summary report with key metrics",
                template_content=EXECUTIVE_SUMMARY_TEMPLATE,
                required_fields=["title", "summary", "metrics"],
                sample_data={
                    "title": "Q4 2024 Report",
                    "summary": "Quarterly performance overview",
                    "metrics": [
                        {"name": "Revenue", "value": "$1.2M", "change": "+15%"},
                    ],
                },
            ),
            "data_analysis": ReportTemplate(
                template_id="data_analysis",
                name="Data Analysis Report",
                description="Detailed data analysis with tables and charts",
                template_content=DATA_ANALYSIS_TEMPLATE,
                required_fields=["title", "data", "analysis"],
                sample_data={
                    "title": "Sales Analysis",
                    "data": [{"product": "A", "sales": 100}],
                    "analysis": "Key findings...",
                },
            ),
            "comparison": ReportTemplate(
                template_id="comparison",
                name="Comparison Report",
                description="Compare two datasets or periods",
                template_content=COMPARISON_TEMPLATE,
                required_fields=["title", "baseline", "comparison"],
                sample_data={
                    "title": "Year over Year",
                    "baseline": {"period": "2023", "value": 100},
                    "comparison": {"period": "2024", "value": 120},
                },
            ),
        }

    @staticmethod
    def _format_datetime(value: Any, fmt: str = "%Y-%m-%d %H:%M") -> str:
        """Format datetime value."""
        if isinstance(value, datetime):
            return value.strftime(fmt)
        if isinstance(value, str):
            try:
                dt = datetime.fromisoformat(value)
                return dt.strftime(fmt)
            except ValueError:
                return value
        return str(value)

    @staticmethod
    def _format_currency(value: Any, symbol: str = "$") -> str:
        """Format currency value."""
        try:
            num = float(value)
            return f"{symbol}{num:,.2f}"
        except (ValueError, TypeError):
            return str(value)

    @staticmethod
    def _format_percentage(value: Any, decimals: int = 1) -> str:
        """Format percentage value."""
        try:
            num = float(value)
            return f"{num:.{decimals}f}%"
        except (ValueError, TypeError):
            return str(value)

    def list_templates(self) -> list[dict[str, Any]]:
        """
        List available report templates.

        Returns:
            List of template metadata dicts.
        """
        templates = []

        # Built-in templates
        for template in self._builtin_templates.values():
            templates.append({
                "template_id": template.template_id,
                "name": template.name,
                "description": template.description,
                "required_fields": template.required_fields,
                "type": "builtin",
            })

        # Custom templates from filesystem
        template_path = Path(self._settings.template_dir)
        if template_path.exists():
            for file in template_path.glob("*.html"):
                templates.append({
                    "template_id": file.stem,
                    "name": file.stem.replace("_", " ").title(),
                    "description": f"Custom template: {file.name}",
                    "required_fields": [],
                    "type": "custom",
                })

        return templates

    async def generate_report(
        self,
        template_id: str,
        data: dict[str, Any],
        output_format: str = "html",
        filename: str | None = None,
    ) -> GeneratedReport:
        """
        Generate a report from a template.

        Args:
            template_id: Template identifier.
            data: Data to populate the template.
            output_format: Output format (html, pdf, docx).
            filename: Optional custom filename.

        Returns:
            GeneratedReport with file details.
        """
        import uuid

        # Get template
        if template_id in self._builtin_templates:
            template = self._builtin_templates[template_id]
            jinja_template = self._jinja_env.from_string(template.template_content)
        else:
            try:
                jinja_template = self._jinja_env.get_template(f"{template_id}.html")
            except Exception as e:
                raise ValueError(f"Template not found: {template_id}") from e

        # Add default context
        context = {
            "company_name": self._settings.company_name,
            "generated_at": datetime.now().isoformat(),
            "report_id": uuid.uuid4().hex[:8],
            **data,
        }

        # Render HTML
        html_content = jinja_template.render(**context)

        # Generate filename
        report_id = uuid.uuid4().hex
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{template_id}_{timestamp}"

        # Generate output
        if output_format == "html":
            output_path = self._save_html(html_content, filename)
        elif output_format == "pdf":
            output_path = await self._save_pdf(html_content, filename)
        elif output_format == "docx":
            output_path = await self._save_docx(html_content, filename, data)
        else:
            raise ValueError(f"Unsupported format: {output_format}")

        file_size = os.path.getsize(output_path) if os.path.exists(output_path) else 0

        return GeneratedReport(
            report_id=report_id,
            template_id=template_id,
            format=output_format,
            filename=os.path.basename(output_path),
            file_path=output_path,
            file_size=file_size,
            generated_at=datetime.now().isoformat(),
            download_url=f"/api/reports/download/{os.path.basename(output_path)}",
        )

    def _save_html(self, html_content: str, filename: str) -> str:
        """Save HTML report."""
        output_path = os.path.join(
            self._settings.output_dir,
            f"{filename}.html",
        )
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        return output_path

    async def _save_pdf(self, html_content: str, filename: str) -> str:
        """Save PDF report using weasyprint or fallback."""
        output_path = os.path.join(
            self._settings.output_dir,
            f"{filename}.pdf",
        )

        try:
            from weasyprint import HTML

            HTML(string=html_content).write_pdf(output_path)
        except ImportError:
            # Fallback: save as HTML with PDF extension note
            logger.warning("weasyprint not installed, saving as HTML")
            html_path = output_path.replace(".pdf", ".html")
            with open(html_path, "w", encoding="utf-8") as f:
                f.write(html_content)
            return html_path

        return output_path

    async def _save_docx(
        self,
        html_content: str,
        filename: str,
        data: dict[str, Any],
    ) -> str:
        """Save DOCX report."""
        from docx import Document
        from docx.shared import Inches, Pt

        output_path = os.path.join(
            self._settings.output_dir,
            f"{filename}.docx",
        )

        doc = Document()

        # Add title
        title = data.get("title", "Report")
        doc.add_heading(title, 0)

        # Add generation info
        doc.add_paragraph(
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        doc.add_paragraph(f"Company: {self._settings.company_name}")

        # Add summary if present
        if "summary" in data:
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(data["summary"])

        # Add metrics if present
        if "metrics" in data and isinstance(data["metrics"], list):
            doc.add_heading("Key Metrics", level=1)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Metric"
            hdr_cells[1].text = "Value"
            hdr_cells[2].text = "Change"

            for metric in data["metrics"]:
                row_cells = table.add_row().cells
                row_cells[0].text = str(metric.get("name", ""))
                row_cells[1].text = str(metric.get("value", ""))
                row_cells[2].text = str(metric.get("change", ""))

        # Add analysis if present
        if "analysis" in data:
            doc.add_heading("Analysis", level=1)
            doc.add_paragraph(data["analysis"])

        # Add data table if present
        if "data" in data and isinstance(data["data"], list) and data["data"]:
            doc.add_heading("Data", level=1)
            first_row = data["data"][0]
            if isinstance(first_row, dict):
                cols = list(first_row.keys())
                table = doc.add_table(rows=1, cols=len(cols))
                table.style = "Table Grid"

                # Header
                for i, col in enumerate(cols):
                    table.rows[0].cells[i].text = str(col)

                # Data rows
                for row_data in data["data"][:50]:  # Limit rows
                    row_cells = table.add_row().cells
                    for i, col in enumerate(cols):
                        row_cells[i].text = str(row_data.get(col, ""))

        doc.save(output_path)
        return output_path

    def create_chart(self, config: ChartConfig) -> str:
        """
        Create a chart and return as base64 image.

        Args:
            config: Chart configuration.

        Returns:
            Base64-encoded PNG image.
        """
        try:
            import matplotlib.pyplot as plt
            import matplotlib

            matplotlib.use("Agg")  # Non-interactive backend

            fig, ax = plt.subplots(figsize=(config.width / 100, config.height / 100))

            if config.chart_type == "bar":
                x = config.data.get("x", [])
                y = config.data.get("y", [])
                ax.bar(x, y)
            elif config.chart_type == "line":
                x = config.data.get("x", [])
                y = config.data.get("y", [])
                ax.plot(x, y, marker="o")
            elif config.chart_type == "pie":
                labels = config.data.get("labels", [])
                values = config.data.get("values", [])
                ax.pie(values, labels=labels, autopct="%1.1f%%")
            elif config.chart_type == "scatter":
                x = config.data.get("x", [])
                y = config.data.get("y", [])
                ax.scatter(x, y)

            ax.set_title(config.title)
            plt.tight_layout()

            # Save to bytes
            buf = io.BytesIO()
            plt.savefig(buf, format="png", dpi=100)
            buf.seek(0)
            plt.close(fig)

            return base64.b64encode(buf.read()).decode("utf-8")

        except Exception as e:
            logger.error(f"Chart creation failed: {e}")
            return ""


# Built-in template content
EXECUTIVE_SUMMARY_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
    <title>{{ title }}</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; }
        .header { border-bottom: 2px solid #333; padding-bottom: 20px; }
        .company { color: #666; font-size: 14px; }
        .title { font-size: 28px; margin: 10px 0; }
        .generated { color: #999; font-size: 12px; }
        .summary { margin: 30px 0; line-height: 1.6; }
        .metrics { display: flex; flex-wrap: wrap; gap: 20px; margin: 30px 0; }
        .metric { background: #f5f5f5; padding: 20px; border-radius: 8px; min-width: 200px; }
        .metric-name { color: #666; font-size: 14px; }
        .metric-value { font-size: 24px; font-weight: bold; margin: 5px 0; }
        .metric-change { font-size: 14px; }
        .positive { color: #28a745; }
        .negative { color: #dc3545; }
    </style>
</head>
<body>
    <div class="header">
        <div class="company">{{ company_name }}</div>
        <h1 class="title">{{ title }}</h1>
        <div class="generated">Generated: {{ generated_at | datetime }}</div>
    </div>

    {% if summary %}
    <div class="summary">
        <h2>Executive Summary</h2>
        <p>{{ summary }}</p>
    </div>
    {% endif %}

    {% if metrics %}
    <div class="metrics">
        {% for metric in metrics %}
        <div class="metric">
            <div class="metric-name">{{ metric.name }}</div>
            <div class="metric-value">{{ metric.value }}</div>
            {% if metric.change %}
            <div class="metric-change {{ 'positive' if '+' in metric.change else 'negative' }}">
                {{ metric.change }}
            </div>
            {% endif %}
        </div>
        {% endfor %}
    </div>
    {% endif %}

    {% if key_points %}
    <div class="key-points">
        <h2>Key Points</h2>
        <ul>
        {% for point in key_points %}
            <li>{{ point }}</li>
        {% endfor %}
        </ul>
    </div>
    {% endif %}
</body>
</html>
"""

DATA_ANALYSIS_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
    <title>{{ title }}</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; }
        table { border-collapse: collapse; width: 100%; margin: 20px 0; }
        th, td { border: 1px solid #ddd; padding: 12px; text-align: left; }
        th { background-color: #f5f5f5; }
        tr:nth-child(even) { background-color: #fafafa; }
        .analysis { margin: 30px 0; line-height: 1.6; }
    </style>
</head>
<body>
    <h1>{{ title }}</h1>
    <p>Generated: {{ generated_at | datetime }}</p>

    {% if data %}
    <h2>Data</h2>
    <table>
        <thead>
            <tr>
            {% for key in data[0].keys() %}
                <th>{{ key }}</th>
            {% endfor %}
            </tr>
        </thead>
        <tbody>
        {% for row in data %}
            <tr>
            {% for value in row.values() %}
                <td>{{ value }}</td>
            {% endfor %}
            </tr>
        {% endfor %}
        </tbody>
    </table>
    {% endif %}

    {% if analysis %}
    <div class="analysis">
        <h2>Analysis</h2>
        <p>{{ analysis }}</p>
    </div>
    {% endif %}
</body>
</html>
"""

COMPARISON_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
    <title>{{ title }}</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; }
        .comparison { display: flex; gap: 40px; margin: 30px 0; }
        .period { flex: 1; padding: 20px; border: 1px solid #ddd; border-radius: 8px; }
        .period-title { font-size: 18px; color: #666; }
        .period-value { font-size: 32px; font-weight: bold; margin: 10px 0; }
        .change { font-size: 24px; text-align: center; margin: 20px 0; }
        .positive { color: #28a745; }
        .negative { color: #dc3545; }
    </style>
</head>
<body>
    <h1>{{ title }}</h1>
    <p>Generated: {{ generated_at | datetime }}</p>

    <div class="comparison">
        <div class="period">
            <div class="period-title">{{ baseline.period }}</div>
            <div class="period-value">{{ baseline.value }}</div>
        </div>
        <div class="period">
            <div class="period-title">{{ comparison.period }}</div>
            <div class="period-value">{{ comparison.value }}</div>
        </div>
    </div>

    {% if change_percentage %}
    <div class="change {{ 'positive' if change_percentage > 0 else 'negative' }}">
        {{ '+' if change_percentage > 0 else '' }}{{ change_percentage }}%
    </div>
    {% endif %}
</body>
</html>
"""


# Singleton instance
_report_service: ReportService | None = None


def get_report_service() -> ReportService:
    """Get or create the singleton report service."""
    global _report_service
    if _report_service is None:
        _report_service = ReportService()
    return _report_service
